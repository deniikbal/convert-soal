import streamlit as st
import docx
from docx import Document
from docx.shared import Cm, Inches

# Antarmuka Streamlit
st.title('Aplikasi Format Soal')
st.markdown("""
Upload file Word yang berisi soal-soal untuk diformat ke dalam template tabel.
File hasil konversi akan tersedia untuk diunduh.
""")

# Fungsi untuk memproses file input
def process_input_file(input_file):
    # Baca file Word
    doc = Document(input_file)
    
    # Ekstrak data soal
    questions = []
    current_question = {}
    question_patterns = ('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')
    option_patterns = ('A.', 'B.', 'C.', 'D.', 'E.')
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Deteksi nomor soal
        if any(text.startswith(p) for p in question_patterns):
            if current_question:
                # Validasi soal lengkap sebelum ditambahkan
                if all(key in current_question for key in ['number', 'question', 'options', 'answer']):
                    questions.append(current_question)
                else:
                    st.warning(f"Soal {current_question.get('number', '')} tidak lengkap dan akan diabaikan")
            
            # Mulai soal baru
            parts = text.split('.', 1)
            if len(parts) > 1:
                current_question = {
                    'number': parts[0].strip(),
                    'question': parts[1].strip(),
                    'options': [],
                    'answer': ''
                }
        
        # Deteksi pilihan jawaban
        elif any(text.startswith(p) for p in option_patterns):
            if current_question:
                current_question['options'].append(text)
        
        # Deteksi kunci jawaban
        elif text.lower().replace(' ', '').startswith(('kunci:', 'jawaban:')):
            if current_question:
                parts = text.split(':', 1)
                if len(parts) > 1:
                    current_question['answer'] = parts[1].strip()
    
    # Tambahkan soal terakhir jika valid
    if current_question and all(key in current_question for key in ['number', 'question', 'options', 'answer']):
        questions.append(current_question)
    
    # Validasi minimal 1 soal
    if not questions:
        st.error('Format file tidak sesuai. Pastikan file berisi:')
        st.markdown("""
        - Nomor soal (contoh: 1. Pertanyaan...)
        - Pilihan jawaban (contoh: A. Jawaban A)
        - Kunci jawaban (contoh: Kunci: A)
        """)
    
    return questions

# Fungsi untuk membuat template output
def create_template(questions):
    # Buat dokumen baru
    doc = Document()
    
    # Pengaturan ukuran kertas A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # Lebar kertas A4 dalam inches
    section.page_height = Inches(11.69)  # Tinggi kertas A4 dalam inches
    section.left_margin = Inches(1)  # Margin kiri dalam inches
    section.right_margin = Inches(1)  # Margin kanan dalam inches
    
    # Iterasi untuk setiap soal
    for i, question in enumerate(questions):
        # Tambahkan jarak antar tabel jika bukan soal pertama
        if i > 0:
            doc.add_paragraph()  # Tambah baris kosong antara tabel
        
        # Buat tabel 2 kolom x 10 baris
        table = doc.add_table(rows=10, cols=2)
        
        # Set lebar kolom
        col1_width = Cm(1.5)  # Lebar kolom 1: 1.5 cm
        col2_width = section.page_width - col1_width - section.left_margin - section.right_margin
        widths = (col1_width, col2_width)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        # Isi template sesuai format baru
        table.cell(0, 0).text = 'TS'
        table.cell(0, 1).text = 'PG'
        table.cell(1, 0).text = 'KD'
        table.cell(1, 1).text = ''  # Nomor soal dihapus dari kolom 2 baris 2
        table.cell(2, 0).text = 'KJ'
        table.cell(2, 1).text = question['answer']
        table.cell(3, 0).text = 'ABS'
        table.cell(3, 1).text = ''  # Kolom 2 baris 2 dikosongkan
        
        # Baris 5: Nomor soal dan teks soal
        table.cell(4, 0).text = question['number']
        table.cell(4, 1).text = question['question']
        
        # Baris 6-10: Pilihan jawaban
        for j, option in enumerate(question['options']):
            # Hapus prefix A., B., C., D., E. dari teks option
            option_text = option.split('.', 1)[1].strip() if '.' in option else option
            table.cell(5 + j, 0).text = ['A', 'B', 'C', 'D', 'E'][j]
            table.cell(5 + j, 1).text = option_text
            
        # Format tabel dengan styling
        table.style = 'Table Grid'  # Tambahkan border ke seluruh tabel
    
    return doc

# Upload file interface
uploaded_file = st.file_uploader("Upload file soal", type=['docx'])

if uploaded_file is not None:
    output_path = 'output.docx'
    try:
        questions = process_input_file(uploaded_file)
        
        if not questions:
            st.error('File tidak berisi soal yang valid. Pastikan format sesuai!')
        else:
            output_doc = create_template(questions)
            output_doc.save(output_path)
            
            # Opsi download file
            with open(output_path, "rb") as file:
                btn = st.download_button(
                    label="Download file output",
                    data=file,
                    file_name="output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                    
    except Exception as e:
        st.error(f'Terjadi error: {str(e)}')

# Tampilkan panduan format file
with st.expander("Format File Input"):
    st.markdown("""
    File Word harus mengikuti format berikut:
    ```
    1. Pertanyaan pertama
    A. Pilihan A
    B. Pilihan B
    C. Pilihan C
    D. Pilihan D
    E. Pilihan E
    Kunci: A

    2. Pertanyaan kedua
    ...
    ```
    """)
